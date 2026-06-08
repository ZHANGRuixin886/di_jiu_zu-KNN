from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = ROOT / "deliverables" / "机器学习大作业报告_KNN手写数字识别.md"
REPORT_DOCX = ROOT / "deliverables" / "机器学习大作业报告_KNN手写数字识别.docx"


def set_font(run, name="Microsoft YaHei", size=10.5, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_font(run, size=9.5, bold=bold)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "B7C2CC")
        borders.append(tag)
    tbl_pr.append(borders)


def add_table_from_markdown(doc, lines, start_index):
    rows = []
    idx = start_index
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        line = lines[idx].strip()
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if not all(set(cell.replace(":", "").replace("-", "").strip()) == set() for cell in cells):
            rows.append(cells)
        idx += 1

    if not rows:
        return idx

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = True
    set_table_borders(table)
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell_text(table.cell(r, c), value, bold=(r == 0))
    doc.add_paragraph()
    return idx


def add_code_block(doc, code_lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.22)
    para.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(code_lines):
        if i > 0:
            para.add_run("\n")
        run = para.add_run(line)
        set_font(run, name="Consolas", size=9)


def build_docx():
    text = REPORT_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        if stripped.startswith("|"):
            i = add_table_from_markdown(doc, lines, i)
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:])
            set_font(run, size=18, bold=True)
            p.paragraph_format.space_after = Pt(14)
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            set_font(run, size=14, bold=True)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            set_font(run, size=12, bold=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif stripped[0:2].isdigit() and stripped[2:4] == ". ":
            p = doc.add_paragraph(style=None)
            run = p.add_run(stripped)
            set_font(run, size=10.5)
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_after = Pt(3)
        else:
            p = doc.add_paragraph()
            run = p.add_run(stripped.replace("`", ""))
            set_font(run, size=10.5)
            p.paragraph_format.first_line_indent = Inches(0.28)
            p.paragraph_format.line_spacing = 1.18
            p.paragraph_format.space_after = Pt(4)
        i += 1

    doc.save(REPORT_DOCX)
    print(REPORT_DOCX)


if __name__ == "__main__":
    build_docx()

