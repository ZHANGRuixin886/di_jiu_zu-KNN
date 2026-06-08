from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "deliverables" / "机器学习大作业报告_KNN手写数字识别.docx"
BACKUP = ROOT / "deliverables" / "机器学习大作业报告_KNN手写数字识别_格式调整前备份.docx"

TABLE_CAPTIONS = [
    "表1 主要代码文件及功能",
    "表2 数据集划分",
    "表3 不同 K 值验证集准确率",
    "表4 测试集各类别分类指标",
    "表5 主要误分类情况",
]

LEVEL1_HEADINGS = {
    "1. 项目背景与任务目标",
    "2. 数据集介绍与预处理",
    "3. KNN 算法原理",
    "4. 代码实现与项目结构",
    "5. 实验设计",
    "6. 实验结果与分析",
    "7. 交互式演示界面",
    "8. 学习体会",
    "9. 总结与展望",
}


def set_run_font(run, east_asia: str, latin: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph_font(paragraph, east_asia: str, latin: str, size_pt: float, bold: bool = False) -> None:
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.text = run.text.replace("`", "")
        set_run_font(run, east_asia, latin, size_pt, bold)


def set_exact_line_spacing(paragraph, points: float) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(points)


def set_single_line_spacing(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1


def normalize_heading_spaces(text: str) -> str:
    match = re.match(r"^(\d+(?:\.\d+)+)\s*(.+)$", text)
    if not match:
        return text
    return f"{match.group(1)}  {match.group(2)}"


def classify_paragraph(text: str) -> str:
    stripped = text.strip()
    if not stripped:
        return "empty"
    if stripped == "基于 KNN 的 MNIST 手写数字识别实验报告":
        return "document_title"
    if stripped == "摘要":
        return "level1"
    if stripped in LEVEL1_HEADINGS:
        return "level1"
    if re.match(r"^\d+(?:\.\d+)+\s+", stripped):
        return "level2"
    if stripped.startswith("表") and re.match(r"^表\d+\s", stripped):
        return "table_caption"
    return "body"


def format_paragraph(paragraph) -> None:
    text = paragraph.text.strip()
    kind = classify_paragraph(text)

    if kind == "empty":
        return

    if kind == "document_title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(12)
        set_single_line_spacing(paragraph)
        set_paragraph_font(paragraph, "黑体", "Times New Roman", 16, True)
        return

    if kind == "level1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(7)
        paragraph.paragraph_format.space_after = Pt(7)
        set_single_line_spacing(paragraph)
        set_paragraph_font(paragraph, "黑体", "Times New Roman", 14, True)
        return

    if kind == "level2":
        normalized = normalize_heading_spaces(text)
        if paragraph.text != normalized:
            paragraph.clear()
            paragraph.add_run(normalized)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(6)
        set_single_line_spacing(paragraph)
        set_paragraph_font(paragraph, "黑体", "Times New Roman", 12, True)
        return

    if kind == "table_caption":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        set_single_line_spacing(paragraph)
        set_paragraph_font(paragraph, "宋体", "Times New Roman", 10.5, False)
        return

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    set_exact_line_spacing(paragraph, 20)
    set_paragraph_font(paragraph, "宋体", "Times New Roman", 12, False)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")
        borders.append(tag)
    tbl_pr.append(borders)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index == 0:
                set_cell_shading(cell, "EDEDED")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                set_single_line_spacing(paragraph)
                set_paragraph_font(paragraph, "宋体", "Times New Roman", 10.5, row_index == 0)


def insert_paragraph_before_table(table, text: str):
    paragraph = table._parent.add_paragraph()
    paragraph.add_run(text)
    table._element.addprevious(paragraph._p)
    return paragraph


def remove_existing_table_captions(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if re.match(r"^表\d+\s", text):
            paragraph._element.getparent().remove(paragraph._element)


def apply_format() -> None:
    if not REPORT.exists():
        raise FileNotFoundError(REPORT)
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)

    doc = Document(REPORT)

    # Page layout: keep a clean academic page with balanced margins.
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    remove_existing_table_captions(doc)

    for index, table in enumerate(doc.tables):
        caption = TABLE_CAPTIONS[index] if index < len(TABLE_CAPTIONS) else f"表{index + 1} 实验结果"
        caption_paragraph = insert_paragraph_before_table(table, caption)
        format_paragraph(caption_paragraph)
        format_table(table)

    for paragraph in doc.paragraphs:
        format_paragraph(paragraph)

    doc.save(REPORT)
    print(REPORT)
    print(BACKUP)


if __name__ == "__main__":
    apply_format()
